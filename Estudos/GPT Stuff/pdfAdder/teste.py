import pyautogui
import pyperclip
import time
from docx import Document

# Função para copiar o texto visível no navegador
def copy_visible_chat():
    time.sleep(2)
    # Simula 'Ctrl+A' para selecionar tudo
    print("Selecionando todo o texto visível...")
    pyautogui.hotkey('ctrl', 'a')
    time.sleep(2)  # Espera 2 segundos para garantir que o conteúdo foi selecionado

    # Simula 'Ctrl+C' para copiar o conteúdo
    print("Copiando o texto...")
    pyautogui.hotkey('ctrl', 'c')
    time.sleep(2)  # Espera 2 segundos para garantir que o conteúdo foi copiado

    # Usa pyperclip para pegar o conteúdo copiado da área de transferência
    copied_text = pyperclip.paste()
    
    # Mostra o conteúdo copiado para verificar
    print("Conteúdo copiado:\n", copied_text)

    return copied_text

# Função para salvar o conteúdo no arquivo Word
def update_word_document(copied_text, filename="exemploquevaiatualizarsempre.docx"):
    try:
        # Tenta abrir o documento existente
        doc = Document(filename)
    except:
        # Cria um novo documento se o arquivo não existir
        doc = Document()

    # Adiciona um título e o conteúdo da conversa copiado sem estilo
    doc.add_paragraph('Conversa Capturada')  # Apenas adiciona o título sem estilo
    doc.add_paragraph(copied_text)

    # Adiciona uma mensagem indicando continuação
    doc.add_paragraph("Continuação da conversa anterior, por favor, siga o mesmo raciocínio.")

    # Salva o documento
    doc.save(filename)
    print(f"Documento atualizado: {filename}")

# Função principal
def main():
    # Copia o texto da conversa visível no navegador
    chat_text = copy_visible_chat()

    # Atualiza o arquivo Word com o texto copiado
    update_word_document(chat_text)

if __name__ == "__main__":
    main()
